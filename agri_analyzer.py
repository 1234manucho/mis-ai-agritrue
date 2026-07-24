"""AgriTrue multimodal analysis service.

This module keeps all Gemini calls on the server. It supports:
- crop and livestock image screening;
- PDF and scanned-document integrity screening;
- DOCX, CSV and TXT content analysis;
- structured JSON responses validated with Pydantic.

AI output is decision support, not a laboratory diagnosis, veterinary prescription,
or legal authentication of a document.
"""

from __future__ import annotations

import base64
import csv
import io
import os
from pathlib import Path
from typing import Any

from docx import Document
from google import genai
from PIL import Image, ImageOps, UnidentifiedImageError
from pydantic import BaseModel, Field, ValidationError

try:
    import pillow_heif

    pillow_heif.register_heif_opener()
except ImportError:  # HEIC/HEIF support is optional at runtime.
    pillow_heif = None


DEFAULT_MODEL = "gemini-3.6-flash"
MAX_TEXT_CHARACTERS = 160_000
MAX_IMAGE_EDGE = 2048


class AnalyzerError(RuntimeError):
    """Safe, user-facing analyzer error."""


class ImageDiagnosis(BaseModel):
    category: str = Field(
        description="One of: crop, animal, non_agricultural, unclear."
    )
    subject: str = Field(description="Visible crop, plant, animal, or object.")
    condition_name: str = Field(
        description="Most likely condition, disease, pest, deficiency, injury, or healthy status."
    )
    diagnosis_status: str = Field(
        description="One of: healthy, possible_issue, likely_issue, unclear, not_applicable."
    )
    confidence_percent: int = Field(
        description="Estimated visual confidence from 0 to 100."
    )
    urgency: str = Field(
        description="One of: routine, soon, urgent, emergency, unknown."
    )
    summary: str
    visible_signs: list[str]
    likely_causes: list[str]
    immediate_actions: list[str]
    prevention_measures: list[str]
    when_to_escalate: list[str]
    photo_quality_notes: list[str]
    disclaimer: str


class ExtractedField(BaseModel):
    label: str
    value: str


class VerificationCheck(BaseModel):
    check: str
    finding: str
    status: str = Field(description="One of: passed, concern, unknown, not_applicable.")


class DocumentAnalysis(BaseModel):
    document_type: str
    document_title: str
    summary: str
    visible_text_excerpt: str
    extracted_fields: list[ExtractedField]
    integrity_status: str = Field(
        description="One of: low_risk, needs_review, high_risk, unreadable, not_verifiable."
    )
    risk_score_percent: int = Field(
        description="Visual and internal-consistency risk score from 0 to 100."
    )
    verification_checks: list[VerificationCheck]
    red_flags: list[str]
    missing_or_unclear_information: list[str]
    agricultural_findings: list[str]
    recommended_next_steps: list[str]
    disclaimer: str


def _get_client() -> genai.Client:
    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise AnalyzerError(
            "Gemini is not configured. Add GEMINI_API_KEY to your environment variables."
        )
    return genai.Client(api_key=api_key)


def _model_name() -> str:
    return os.getenv("GEMINI_MODEL", DEFAULT_MODEL).strip() or DEFAULT_MODEL


def _bounded_percent(value: int | float | None) -> int:
    try:
        return max(0, min(100, int(round(float(value or 0)))))
    except (TypeError, ValueError):
        return 0


def normalize_image(source_path: str | Path, destination_path: str | Path) -> Path:
    """Rotate, resize and convert a phone image to a web-safe JPEG."""
    source = Path(source_path)
    destination = Path(destination_path)
    destination.parent.mkdir(parents=True, exist_ok=True)

    try:
        with Image.open(source) as image:
            image = ImageOps.exif_transpose(image)
            image.thumbnail((MAX_IMAGE_EDGE, MAX_IMAGE_EDGE), Image.Resampling.LANCZOS)

            if image.mode in {"RGBA", "LA"}:
                background = Image.new("RGB", image.size, "white")
                alpha = image.getchannel("A")
                background.paste(image.convert("RGB"), mask=alpha)
                image = background
            elif image.mode != "RGB":
                image = image.convert("RGB")

            image.save(destination, format="JPEG", quality=88, optimize=True)
    except (UnidentifiedImageError, OSError, ValueError) as exc:
        raise AnalyzerError(
            "The selected image could not be opened. Use a clear JPG, PNG, WEBP, HEIC, or HEIF photo."
        ) from exc

    return destination


def _interaction_json(
    *,
    prompt: str,
    schema: type[BaseModel],
    media_type: str | None = None,
    media_bytes: bytes | None = None,
    media_kind: str | None = None,
) -> BaseModel:
    client = _get_client()
    input_items: list[dict[str, Any]] = [{"type": "text", "text": prompt}]

    if media_bytes is not None and media_type and media_kind:
        input_items.append(
            {
                "type": media_kind,
                "data": base64.b64encode(media_bytes).decode("utf-8"),
                "mime_type": media_type,
            }
        )

    try:
        interaction = client.interactions.create(
            model=_model_name(),
            input=input_items,
            response_format={
                "type": "text",
                "mime_type": "application/json",
                "schema": schema.model_json_schema(),
            },
        )
        parsed = schema.model_validate_json(interaction.output_text)
    except ValidationError as exc:
        raise AnalyzerError("The AI returned an incomplete result. Please try again.") from exc
    except Exception as exc:
        message = str(exc).lower()
        if "api key" in message or "permission" in message or "unauth" in message:
            raise AnalyzerError(
                "The Gemini API key was rejected. Replace or restrict the key, then redeploy."
            ) from exc
        if "quota" in message or "429" in message or "rate" in message:
            raise AnalyzerError(
                "The AI service is temporarily busy or its quota has been reached. Please retry shortly."
            ) from exc
        if "model" in message and ("not found" in message or "shutdown" in message):
            raise AnalyzerError(
                "The configured Gemini model is unavailable. Set GEMINI_MODEL to a current stable model."
            ) from exc
        raise AnalyzerError("The AI analysis service could not complete this request.") from exc

    return parsed


def analyze_agricultural_image(image_path: str | Path) -> ImageDiagnosis:
    path = Path(image_path)
    prompt = """
You are AgriTrue's agricultural visual screening assistant.

Inspect this single image and return only the requested structured result.

Tasks:
1. Classify the main subject as crop, animal, non_agricultural, or unclear.
2. For a crop: identify the crop where possible and screen for visible disease,
   pest damage, nutrient deficiency, water stress, physical injury, or healthy growth.
3. For an animal: identify the animal where possible and screen for visible illness,
   wounds, parasites, swelling, skin/eye problems, poor body condition, or healthy appearance.
4. Explain only signs actually visible in the photo. Do not invent laboratory results.
5. Give practical immediate actions and prevention measures suitable for a farmer.
6. Do not provide veterinary drug dosages or unsafe pesticide mixing instructions.
7. Mark urgency as emergency only for clearly visible life-threatening signs.
8. Lower confidence when the image is blurry, too distant, poorly lit, obstructed,
   or does not show the affected area.
9. If multiple diagnoses are possible, use a broad condition name and list alternatives
   under likely causes rather than claiming certainty.
10. Include a clear disclaimer that confirmation by a qualified agronomist, plant clinic,
    veterinarian, or laboratory may be required.
"""

    result = _interaction_json(
        prompt=prompt,
        schema=ImageDiagnosis,
        media_type="image/jpeg",
        media_bytes=path.read_bytes(),
        media_kind="image",
    )
    assert isinstance(result, ImageDiagnosis)
    result.confidence_percent = _bounded_percent(result.confidence_percent)
    return result


def _read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _read_csv_file(path: Path) -> str:
    text = _read_text_file(path)
    rows: list[str] = []
    try:
        reader = csv.reader(io.StringIO(text))
        for index, row in enumerate(reader):
            rows.append(" | ".join(cell.strip() for cell in row))
            if index >= 4_999:
                rows.append("[CSV truncated after 5,000 rows]")
                break
    except csv.Error:
        return text
    return "\n".join(rows)


def _read_docx_file(path: Path) -> str:
    try:
        document = Document(path)
    except Exception as exc:
        raise AnalyzerError("The Word document is damaged or cannot be opened.") from exc

    chunks: list[str] = []
    chunks.extend(p.text.strip() for p in document.paragraphs if p.text.strip())

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                chunks.append(" | ".join(values))

    return "\n".join(chunks)


def _document_prompt(extra_text: str = "") -> str:
    return f"""
You are AgriTrue's document screening assistant. Analyze the attached document or
its extracted text and return only the requested structured result.

Purpose:
- Extract important visible fields and summarize the document.
- Screen for internal inconsistencies, missing fields, suspicious formatting,
  altered-looking areas, contradictory dates/numbers, incomplete signatures/stamps,
  and other red flags that are actually visible or present in the text.
- If it is an agricultural report, laboratory result, farm record, veterinary note,
  produce record, invoice, certificate, or inspection form, explain the agricultural findings.

Strict rules:
1. Never claim that a document is legally authentic, genuine, forged, or officially verified.
2. The integrity status is only an AI risk screening based on the uploaded content.
3. Set not_verifiable when issuer databases, QR validation, serial-number lookup,
   metadata, or an original copy would be needed.
4. Set unreadable when the content cannot be read reliably.
5. Keep the visible text excerpt brief and do not reproduce the whole document.
6. Recommend confirmation with the issuing institution when authenticity matters.
7. Do not expose sensitive personal numbers in full; mask all but the last four characters.

Extracted text supplied by the server, when available:
---
{extra_text[:MAX_TEXT_CHARACTERS]}
---
"""


def analyze_document_file(file_path: str | Path, extension: str) -> DocumentAnalysis:
    path = Path(file_path)
    extension = extension.lower().lstrip(".")

    if extension == "pdf":
        data = path.read_bytes()
        if not data.startswith(b"%PDF"):
            raise AnalyzerError("The uploaded file is not a valid PDF.")
        result = _interaction_json(
            prompt=_document_prompt(),
            schema=DocumentAnalysis,
            media_type="application/pdf",
            media_bytes=data,
            media_kind="document",
        )
    elif extension in {"jpg", "jpeg", "png", "webp", "heic", "heif"}:
        normalized = path.with_name(f"{path.stem}-normalized.jpg")
        try:
            normalize_image(path, normalized)
            result = _interaction_json(
                prompt=_document_prompt(),
                schema=DocumentAnalysis,
                media_type="image/jpeg",
                media_bytes=normalized.read_bytes(),
                media_kind="image",
            )
        finally:
            normalized.unlink(missing_ok=True)
    elif extension == "docx":
        text = _read_docx_file(path)
        if not text.strip():
            raise AnalyzerError(
                "No readable text was found in the Word document. Scan it as a PDF or image instead."
            )
        result = _interaction_json(
            prompt=_document_prompt(text),
            schema=DocumentAnalysis,
        )
    elif extension == "csv":
        text = _read_csv_file(path)
        result = _interaction_json(
            prompt=_document_prompt(text),
            schema=DocumentAnalysis,
        )
    elif extension == "txt":
        text = _read_text_file(path)
        result = _interaction_json(
            prompt=_document_prompt(text),
            schema=DocumentAnalysis,
        )
    else:
        raise AnalyzerError("Unsupported document format.")

    assert isinstance(result, DocumentAnalysis)
    result.risk_score_percent = _bounded_percent(result.risk_score_percent)
    return result


def generate_farming_chat_reply(message: str) -> str:
    """Small compatibility helper for the app's existing farming chatbot routes."""
    clean_message = (message or "").strip()
    if not clean_message:
        raise AnalyzerError("No message was provided.")

    client = _get_client()
    prompt = f"""
You are AgriTrue, a practical agricultural assistant for African farmers.
Answer clearly and concisely. Distinguish general guidance from advice that
requires a veterinarian, agronomist, laboratory, or local authority.

Farmer's question: {clean_message}
"""
    try:
        interaction = client.interactions.create(
            model=_model_name(),
            input=prompt,
        )
        return interaction.output_text.strip()
    except Exception as exc:
        raise AnalyzerError("The farming assistant is currently unavailable.") from exc